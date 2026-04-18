#!/usr/bin/env python3
"""
draft_proposal.tex içeriğini Proposal Template.docx şablonuna aktarır.
Bölüm başlıkları ve Word stilleri şablondan kalır; talimat metinleri LaTeX içeriğiyle değiştirilir.
"""

from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from bib_loader import resolve_reference_lines
from edit_proposal import save_with_retry
from proposal_tex_import import (
    OBJECTIVES_TASK_ROWS,
    SECTION_ORDER,
    ContentLine,
    parse_proposal_tex,
)

DEFAULT_TEMPLATE = Path(__file__).resolve().parent / "Proposal Template.docx"
DEFAULT_TEX = Path(__file__).resolve().parent / "draft_proposal.tex"
DEFAULT_GANTT = Path(__file__).resolve().parent / "ens491_gantt.png"


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph_index(doc: Document, text: str, start: int = 0) -> int:
    t = text.strip()
    for i in range(start, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == t:
            return i
    return -1


def find_next_heading_index(doc: Document, after_idx: int, headings: list[str]) -> int:
    for i in range(after_idx + 1, len(doc.paragraphs)):
        tx = doc.paragraphs[i].text.strip()
        if tx in headings:
            return i
    return len(doc.paragraphs)


def set_cover_line(paragraph: Paragraph, prefix: str, value: str) -> None:
    runs = paragraph.runs
    line = f"{prefix}{value}"
    if runs:
        runs[0].text = line
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(line)


def apply_cover(doc: Document, cover: dict[str, str]) -> None:
    """Şablondaki kapak satırlarını günceller (indeksler Proposal Template ile uyumlu)."""
    if cover.get("course_line"):
        cover["course_line"] = cover["course_line"].replace("--", "\u2013")
    # 0: kurs, 1: Proposal
    if len(doc.paragraphs) > 0 and cover.get("course_line"):
        p0 = doc.paragraphs[0]
        if p0.runs:
            p0.runs[0].text = cover["course_line"]
            for r in p0.runs[1:]:
                r.text = ""
    if len(doc.paragraphs) > 1 and cover.get("proposal_line"):
        p1 = doc.paragraphs[1]
        if p1.runs:
            p1.runs[0].text = cover["proposal_line"]
            for r in p1.runs[1:]:
                r.text = ""

    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t.lower().startswith("project title"):
            set_cover_line(para, "Project Title: ", cover.get("project_title", ""))
        elif t.lower().startswith("group members"):
            set_cover_line(para, "Group Members: ", cover.get("group_members", ""))
        elif t.lower().startswith("supervisor"):
            set_cover_line(para, "Supervisor(s): ", cover.get("supervisors", ""))
        elif t.strip() == "Date:" or t.lower().startswith("date:"):
            set_cover_line(para, "Date: ", cover.get("date", ""))


def insert_objectives_table_after(
    doc: Document,
    anchor: Paragraph,
    rows: list[tuple[str, str]],
    body_style: str,
) -> Paragraph:
    """Objectives/Tasks için 2 sütunlu Word tablosu ekler; şablon gövdesinde anchor'dan sonra gelir."""
    table = doc.add_table(rows=len(rows), cols=2)
    tbl_el = table._tbl
    body = doc.element.body
    body.remove(tbl_el)
    anchor._element.addnext(tbl_el)

    for i, (left, right) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].text = left
        row.cells[1].text = right
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    spacer = OxmlElement("w:p")
    tbl_el.addnext(spacer)
    return Paragraph(spacer, anchor._parent)


def fill_body_section(
    doc: Document,
    heading: str,
    lines: list[ContentLine],
    all_headings: list[str],
    *,
    body_style_hint: str | None = None,
) -> None:
    hi = find_paragraph_index(doc, heading)
    if hi < 0:
        print(f"Uyarı: şablonda '{heading}' başlığı bulunamadı.", file=sys.stderr)
        return

    ni = find_next_heading_index(doc, hi, all_headings)
    if ni <= hi + 1 and not lines:
        return

    body_style = body_style_hint or "Normal"
    if ni > hi + 1:
        body_style = doc.paragraphs[hi + 1].style.name

    to_remove = [doc.paragraphs[i] for i in range(hi + 1, ni)]
    for p in reversed(to_remove):
        delete_paragraph(p)

    anchor = doc.paragraphs[hi]
    list_style = "List Paragraph"
    for cl in lines:
        if cl.kind == "subheading":
            anchor = insert_paragraph_after(anchor, style=body_style)
            r = anchor.add_run(cl.text)
            r.bold = True
        elif cl.kind == "objectives_table":
            anchor = insert_objectives_table_after(
                doc, anchor, OBJECTIVES_TASK_ROWS, body_style
            )
        elif cl.kind == "list":
            anchor = insert_paragraph_after(anchor, style=list_style)
            anchor.add_run(cl.text)
        else:
            anchor = insert_paragraph_after(anchor, style=body_style)
            anchor.add_run(cl.text)


def insert_gantt_in_schedule(
    doc: Document,
    all_headings: list[str],
    image_path: Path,
    *,
    width_inches: float = 6.2,
    caption: str | None = None,
) -> None:
    """PROJECT SCHEDULE bölümünün sonuna Gantt görseli ekler."""
    heading = "PROJECT SCHEDULE"
    hi = find_paragraph_index(doc, heading)
    if hi < 0:
        print("Uyarı: PROJECT SCHEDULE başlığı yok; Gantt eklenmedi.", file=sys.stderr)
        return
    ni = find_next_heading_index(doc, hi, all_headings)
    last_idx = ni - 1 if ni > hi + 1 else hi
    anchor = doc.paragraphs[last_idx]

    pic_p = insert_paragraph_after(anchor, style="Normal")
    run = pic_p.add_run()
    if image_path.is_file():
        run.add_picture(str(image_path), width=Inches(width_inches))
    else:
        pic_p.add_run(f"[Gantt görseli bulunamadı: {image_path}]")

    if caption:
        cap = insert_paragraph_after(pic_p, style="Normal")
        r = cap.add_run(caption)
        r.italic = True


def main() -> int:
    parser = argparse.ArgumentParser(
        description="draft_proposal.tex → Proposal Template.docx (şablon stillerini korur)",
    )
    parser.add_argument(
        "--tex",
        type=Path,
        default=DEFAULT_TEX,
        help="Kaynak .tex dosyası",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=DEFAULT_TEMPLATE,
        help="Hedef Word şablonu",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Çıktı .docx (varsayılan: --template ile aynı dosyanın üzerine yazar)",
    )
    parser.add_argument(
        "--backup",
        "-b",
        action="store_true",
        help="Üzerine yazmadan önce .bak yedeği al",
    )
    parser.add_argument(
        "--gantt",
        type=Path,
        default=DEFAULT_GANTT,
        help=f"PROJECT SCHEDULE için PNG (varsayılan: {DEFAULT_GANTT.name})",
    )
    parser.add_argument(
        "--no-gantt",
        action="store_true",
        help="Gantt görselini ekleme",
    )
    parser.add_argument(
        "--references-bib",
        type=Path,
        default=None,
        help="references.bib yolu (yoksa \\bibliography{...} ile aynı adda .bib aranır)",
    )
    parser.add_argument(
        "--references-bbl",
        type=Path,
        default=None,
        help="pdflatex/bibtex ile üretilmiş .bbl (varsa .bib yerine tercih edilir)",
    )
    args = parser.parse_args()

    tex_path = args.tex.expanduser().resolve()
    template_path = args.template.expanduser().resolve()
    out_path = (args.output or template_path).expanduser().resolve()

    if not tex_path.is_file():
        print(f".tex bulunamadı: {tex_path}", file=sys.stderr)
        return 1
    if not template_path.is_file():
        print(f"Şablon bulunamadı: {template_path}", file=sys.stderr)
        return 1

    cover, sections = parse_proposal_tex(tex_path)

    ref_lines = resolve_reference_lines(
        tex_path,
        bbl_path=args.references_bbl.expanduser().resolve() if args.references_bbl else None,
        bib_path=args.references_bib.expanduser().resolve() if args.references_bib else None,
    )
    if ref_lines:
        sections["REFERENCES"] = [ContentLine("normal", ln) for ln in ref_lines]

    doc = Document(str(template_path))
    apply_cover(doc, cover)

    # Şablondaki ana bölüm başlıkları (tam metin eşleşmesi)
    all_headings = list(SECTION_ORDER)

    for name in SECTION_ORDER:
        fill_body_section(doc, name, sections.get(name, []), all_headings)

    if not args.no_gantt:
        gantt_path = args.gantt.expanduser().resolve()
        insert_gantt_in_schedule(
            doc,
            all_headings,
            gantt_path,
            caption="Figure: Summary Gantt chart for ENS 491 (task leaders and weeks).",
        )

    if args.backup and out_path.is_file():
        bak = out_path.with_suffix(out_path.suffix + ".bak")
        shutil.copy2(out_path, bak)
        print(f"Yedek: {bak}")

    save_with_retry(out_path, doc)
    print(f"Tamam: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
