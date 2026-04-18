#!/usr/bin/env python3
"""
Proposal_Filled.docx: şablon talimat cümlelerini temizler, APA parantez içi atıf ve
APA 7 kaynakça biçimine geçirir (yapı ve paragraf sayısı korunur).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
DOC_PATH = ROOT / "Proposal_Filled.docx"
GANTT_PNG = ROOT / "ens491_gantt.png"

# Şablondaki ana/alt başlık satırları (Normal stil, tamamı kalın)
_BOLD_PARA_INDICES = frozenset(
    {
        0,
        1,
        7,
        14,
        16,
        19,
        28,
        31,
        38,
        50,
        52,
        61,
        65,
        70,
        81,
        88,
        94,
        104,
    }
)


def _clear_paragraph_children_keep_ppr(p: Paragraph) -> None:
    el = p._element
    for child in list(el):
        if child.tag != qn("w:pPr"):
            el.remove(child)


def insert_gantt_figure(p: Paragraph, image_path: Path, caption: str, *, width_inches: float = 6.3) -> None:
    _clear_paragraph_children_keep_ppr(p)
    if image_path.is_file():
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        p.add_run("\n")
    p.add_run(caption)


def strip_bullets_to_normal(doc: Document) -> None:
    """List Paragraph → Normal; numPr kaldırılır (madde işareti/sayı yok)."""
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if not st.startswith("List"):
            continue
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            num_pr = pPr.find(qn("w:numPr"))
            if num_pr is not None:
                pPr.remove(num_pr)
        p.style = doc.styles["Normal"]


def apply_template_heading_bold(doc: Document) -> None:
    for i in _BOLD_PARA_INDICES:
        if i >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[i]
        if not (p.text or "").strip():
            continue
        for r in p.runs:
            r.bold = True

# Sıralama: soyad (küçük harf), yıl → (parantez içi metin parçası)
_IN = {
    "ranasinghe2019visualising": ("ranasinghe", 2019, "Ranasinghe et al., 2019"),
    "burigat2011pedestrian": ("burigat", 2011, "Burigat & Chittaro, 2011"),
    "mckenzie2016assessing": ("mckenzie", 2016, "McKenzie et al., 2016"),
    "hegarty2016where": ("hegarty", 2016, "Hegarty et al., 2016"),
    "lodha2002visualization": ("lodha", 2002, "Lodha et al., 2002"),
    "Zheng2015": ("zheng", 2015, "Zheng, 2015"),
    "Sobral2019": ("sobral", 2019, "Sobral et al., 2019"),
    "Fadloun2022": ("fadloun", 2022, "Fadloun et al., 2022"),
    "Feng2021": ("feng", 2021, "Feng et al., 2021"),
    "MillardBall2019": ("millard-ball", 2019, "Millard-Ball et al., 2019"),
}


def _replace_bracket_citations(text: str) -> str:
    """Yalnızca BibTeX anahtarı biçimindeki köşeli atıfları değiştirir; LaTeX formülleri korunur."""

    def repl(m: re.Match[str]) -> str:
        inner = m.group(1).strip()
        if not inner:
            return m.group(0)
        if re.fullmatch(r"\d+(,\d+)*", inner.replace(" ", "")):
            return m.group(0)
        keys = [k.strip() for k in inner.split(",") if k.strip()]
        if not keys:
            return m.group(0)
        missing = [k for k in keys if k not in _IN]
        if missing:
            return m.group(0)
        items = sorted((_IN[k] for k in keys), key=lambda t: (t[0], t[1]))
        inner_cite = "; ".join(t[2] for t in items)
        return f"({inner_cite})"

    return re.sub(r"\[([A-Za-z0-9,]+)\]", repl, text)


def _calibri12(doc: Document) -> None:
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(12)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Calibri"
                        r.font.size = Pt(12)


def _fix_block(text: str) -> str:
    return _replace_bracket_citations(text)


APA_REFERENCES_BLOCK = """Burigat, S., & Chittaro, L. (2011). Pedestrian navigation with degraded GPS signal: Investigating the effects of visualizing position uncertainty. In Proceedings of the 13th International Conference on Human Computer Interaction with Mobile Devices and Services (MobileHCI ’11). Association for Computing Machinery. https://doi.org/10.1145/2037373.2037407

Fadloun, S., Morakeb, Y., Cuenca, E., & Choutri, K. (2022). TrajectoryVis: A visual approach to explore movement trajectories. Social Network Analysis and Mining, 12(1), 53. https://doi.org/10.1007/s13278-022-00879-8

Feng, Z., Li, H., Zeng, W., Yang, S.-H., & Qu, H. (2021). Topology density map for urban data visualization and analysis. IEEE Transactions on Visualization and Computer Graphics, 27(2), 828–838. https://doi.org/10.1109/TVCG.2020.3030469

Hegarty, M., Friedman, A., Boone, A. P., & Barrett, T. J. (2016). Where are you? The effect of uncertainty and its visual representation on location judgments in GPS-like displays. Journal of Experimental Psychology: Applied, 22(4), 381–392. https://doi.org/10.1037/xap0000103

Lodha, S. K., Charaniya, A. P., Faaland, N. M., & Ramalingam, S. (2002). Visualization of spatio-temporal GPS uncertainty within a GIS environment. Proceedings of SPIE: Radar Sensor Technology and Data Visualization, 4744. SPIE.

McKenzie, G., Hegarty, M., Barrett, T., & Goodchild, M. (2016). Assessing the effectiveness of different visualizations for judgments of positional uncertainty. International Journal of Geographical Information Science, 30(2), 221–239. https://doi.org/10.1080/13658816.2015.1082566

Millard-Ball, A., Hampshire, R. C., & Weinberger, R. R. (2019). Map-matching poor-quality GPS data in urban environments: The pgMapMatch package. Transportation Planning and Technology, 42(6), 539–553. https://doi.org/10.1080/03081060.2019.1622249

Ranasinghe, C., Schiestel, N., & Kray, C. (2019). Visualising location uncertainty to support navigation under degraded GPS signals: A comparison study. In Proceedings of the 21st International Conference on Human-Computer Interaction with Mobile Devices and Services (MobileHCI ’19). Association for Computing Machinery. https://doi.org/10.1145/3338286.3340128

Sobral, T., Galvão, T., & Borges, J. (2019). Visualization of urban mobility data from intelligent transportation systems. Sensors, 19(2), Article 332. https://doi.org/10.3390/s19020332

Zheng, Y. (2015). Trajectory data mining: An overview. ACM Transactions on Intelligent Systems and Technology, 6(3), Article 1. https://doi.org/10.1145/2743025"""


def main() -> None:
    doc = Document(str(DOC_PATH))

    # Şablon talimatları → boş veya öz metin (paragraf sayısı korunur)
    doc.paragraphs[24].text = ""
    doc.paragraphs[25].text = ""
    doc.paragraphs[26].text = ""

    doc.paragraphs[51].text = (
        "Table 1 summarizes the objectives and tasks of the project and the corresponding intended results."
    )

    p53 = doc.paragraphs[53].text
    if p53.strip().startswith("Describe here"):
        idx = p53.find(":")
        rest = p53[idx + 1 :].strip() if idx != -1 else p53.strip()
        doc.paragraphs[53].text = "We address realistic constraints as follows. " + rest[:1].upper() + rest[1:]

    p66 = doc.paragraphs[66].text
    if "Evaluate environmental impact" in p66 or "simplified LCA mindset" in p66:
        rest = p66.split(". ", 1)[1] if ". " in p66 else p66
        doc.paragraphs[66].text = (
            "We assess environmental implications with a simplified life-cycle view (not a full ISO-compliant inventory). "
            + rest
        )

    doc.paragraphs[72].text = (
        "The life-cycle cost review considers the following categories, where applicable:"
    )

    doc.paragraphs[89].text = (
        "Figure 1 is a Gantt chart for the ENS 491 phase; it summarizes weekly milestones, dependencies, and task leadership."
    )
    doc.paragraphs[90].text = (
        "Each objective has a planned start period and duration aligned with term milestones; implementation and evaluation "
        "tasks are sequenced so that data collection precedes prototype hardening and reporting."
    )
    doc.paragraphs[91].text = (
        "Each task has a single task leader; other members contribute, but accountability is clear for scheduling and deliverables."
    )
    doc.paragraphs[92].text = (
        "If tasks slip, we prioritize a demonstrable prototype and defer large-scale user studies to ENS 492; we mitigate "
        "delays by narrowing encoding comparisons, using public trajectory datasets, and re-scoping optional analyses."
    )

    doc.paragraphs[102].text = (
        "We retain explicit discussion of ethics because location and uncertainty communication can affect privacy and trust "
        "even when no formal human-subjects protocol is required for early prototyping."
    )

    doc.paragraphs[105].text = APA_REFERENCES_BLOCK

    # Yinelenen “Yazar et al. fiil … (Yazar et al., yıl)” → anlatı biçimi “Yazar et al. (yıl) fiil …”
    doc.paragraphs[34].text = (
        "Navigation under degraded GPS. Ranasinghe et al. (2019) evaluated novel uncertainty visualizations for "
        "pedestrian navigation in the field; landmark-oriented designs helped users judge true location under poor GPS. "
        "Burigat and Chittaro (2011) compared circular uncertainty regions and street-based encodings; coloring streets "
        "reduced workload and was rated helpful. These studies show that encoding choice changes performance but often "
        "focus on specific navigation tasks rather than general decision dashboards. Blue-dot and judgment tasks. "
        "McKenzie et al. (2016) compared uniform circles, Gaussian fades, and centroid dots for positional uncertainty "
        "judgments; uniform circles yielded faster, distribution-consistent judgments in their setting. Hegarty et al. "
        "(2016) showed that users combine uncertain cues in task-dependent ways. Shortcoming: uniform disks do not "
        "communicate where probability mass is higher inside the disk—a gap our proposed boundary- and interior-density "
        "ideas target. GIS and trajectories. Lodha et al. (2002) modeled horizontal, speed, and direction error for "
        "moving GPS objects and visualized diffused representations over imagery. Broader trajectory and mobility "
        "literature emphasizes preprocessing (filtering, segmentation, map matching), clutter, and density visualization "
        "(Fadloun et al., 2022; Feng et al., 2021; Millard-Ball et al., 2019; Sobral et al., 2019; Zheng, 2015). "
        "Shortcoming: classical density maps can hide individual structure; planar KDE may misrepresent "
        "network-constrained motion. Commercial products. Consumer maps use proprietary uncertainty metaphors; details "
        "are not open for replication. Our work emphasizes transparent, research-grade encodings we can evaluate and publish."
    )

    # Köşeli parantez atıfları → APA (93: Gantt görseli var; .text silinmesin)
    for i, p in enumerate(doc.paragraphs):
        if i == 93:
            continue
        p.text = _fix_block(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = _fix_block(p.text)

    strip_bullets_to_normal(doc)

    insert_gantt_figure(
        doc.paragraphs[93],
        GANTT_PNG,
        "Figure 1. ENS 491 draft schedule (indicative weeks, task leaders: Mustafa Bozyel, Ömer Mert Özel, Senih Kırmaç).",
        width_inches=6.3,
    )

    _calibri12(doc)
    apply_template_heading_bold(doc)
    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    main()
