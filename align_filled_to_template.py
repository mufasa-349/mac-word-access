#!/usr/bin/env python3
"""
Proposal_Filled.docx kaynak içeriğini Proposal Template.docx iskeletine yerleştirir:
aynı paragraf sayısı ve Word stilleri; gövde metni Calibri 12.
"""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "Proposal Template.docx"
FILLED_BACKUP = ROOT / "_Proposal_Filled_source.docx"
OUT = ROOT / "Proposal_Filled.docx"
GANTT = ROOT / "ens491_gantt.png"


def calibri12_doc(doc: Document) -> None:
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(12)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Calibri"
                        r.font.size = Pt(12)


def copy_table_after(src_doc: Document, dst_doc: Document, after_para_idx: int) -> None:
    if not src_doc.tables:
        return
    clone = deepcopy(src_doc.tables[0]._tbl)
    anchor = dst_doc.paragraphs[after_para_idx]._element
    anchor.addnext(clone)


def insert_gantt(paragraph: Paragraph, image_path: Path, caption: str) -> None:
    paragraph.text = ""
    if image_path.is_file():
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6))
        paragraph.add_run("\n")
    paragraph.add_run(caption)


def fix_tex_escapes(s: str) -> str:
    return s.replace("\\$", "$").replace("\\\\", "\\")


def main() -> None:
    if not TEMPLATE.is_file():
        raise SystemExit(f"Şablon bulunamadı: {TEMPLATE}")

    if not OUT.is_file():
        raise SystemExit(f"Kaynak bulunamadı: {OUT}")

    shutil.copy2(OUT, FILLED_BACKUP)
    src = Document(str(FILLED_BACKUP))

    shutil.copy2(TEMPLATE, OUT)
    doc = Document(str(OUT))

    # Kapak
    for i in (7, 14, 16, 19):
        doc.paragraphs[i].text = src.paragraphs[i].text

    # ABSTRACT
    doc.paragraphs[29].text = (src.paragraphs[29].text + " " + src.paragraphs[30].text).strip()

    # INTRODUCTION (şablonda 32 + 4 madde)
    intro_overview = (
        "The official project theme is AI-Driven Uncertainty-Aware Visualization for Decision Making Under Stress. "
        "Software frequently masks uncertainty, which is especially risky when operators must act quickly. "
        "We scope the first track to GPS and location uncertainty visualization: a reported position should be shown as "
        "a distribution or region, not a single pixel. Location uncertainty arises from sensor noise, multipath, latency, "
        "map-matching and aggregation errors, and model-based inference; we distinguish positional, semantic, model, "
        "and temporal uncertainty to structure the design space."
    )
    doc.paragraphs[32].text = intro_overview

    doc.paragraphs[33].text = " ".join(src.paragraphs[i].text for i in (33, 34, 35)).strip()
    doc.paragraphs[34].text = " ".join(src.paragraphs[i].text for i in (37, 38, 39, 40)).strip()
    doc.paragraphs[35].text = src.paragraphs[42].text.strip()
    doc.paragraphs[36].text = src.paragraphs[44].text.strip()

    # PROPOSED SOLUTION AND METHODS
    doc.paragraphs[39].text = (
        "This part introduces our approach: an iterative pipeline spanning literature review, encoding design for GPS "
        "uncertainty, interactive prototyping (web tool with point, heatmap, kernel-density, and accuracy-ring modes on "
        "real Fix logs), and evaluation of interpretability and honest communication—toward later stress-aware decision support."
    )
    doc.paragraphs[40].text = src.paragraphs[47].text
    doc.paragraphs[41].text = (
        src.paragraphs[48].text + " " + src.paragraphs[49].text + " " + src.paragraphs[50].text
    ).strip()

    doc.paragraphs[42].text = (
        "Explain why this project solves a complex problem: uncertainty must be communicated under perceptual limits, "
        "readability, latency, and trust; no single glyph fits all tasks; the work combines HCI, visualization, GIS, "
        "and trajectory research; integration of ingestion, UI, and evaluation is non-routine; misleading confidence "
        "has safety and privacy consequences; the system spans data, UI, evaluation, and future AI layers."
    )
    for i, j in zip(range(43, 49), range(52, 58)):
        doc.paragraphs[i].text = src.paragraphs[j].text

    # Objectives tablosu: şablonda 51. paragraftan sonra
    copy_table_after(src, doc, 51)

    # Realistic Constraints
    doc.paragraphs[53].text = (
        "Describe here the realistic constraints that affected this project and how we addressed them: we rely on "
        "student-owned devices and open map tiles; we minimize travel for logging; we treat location traces as sensitive data."
    )
    doc.paragraphs[54].text = (
        "Economic: No dedicated hardware budget; we use student devices and free/open tiles (e.g., OpenStreetMap); "
        "engineering effort is the main cost driver."
    )
    doc.paragraphs[55].text = (
        "Environmental: Field logging has negligible environmental impact; we avoid unnecessary travel for data collection."
    )
    doc.paragraphs[56].text = src.paragraphs[62].text
    doc.paragraphs[57].text = src.paragraphs[63].text
    doc.paragraphs[58].text = src.paragraphs[64].text
    doc.paragraphs[59].text = (
        "Sustainability (social, economic, and environmental): Open text logs support reproducibility; documentation "
        "reduces maintainer lock-in; static hosting keeps operational footprint low."
    )

    # Engineering / Scientific Standards
    eng = src.paragraphs[66].text
    mid = max(len(eng) // 2, 1)
    split_at = eng.rfind(". ", 0, mid)
    if split_at == -1:
        split_at = eng.find(". ")
    if split_at != -1:
        a, b = eng[: split_at + 1].strip(), eng[split_at + 1 :].strip()
    else:
        a, b = eng, ""
    doc.paragraphs[62].text = a
    doc.paragraphs[63].text = b or (
        "A technical or engineering standard is an established norm used to ensure quality, reliability, and safety in empirical studies."
    )
    doc.paragraphs[64].text = (
        "A code is a set of rules followed in design or inspection; our software prototype follows documented parameters, "
        "honest labeling, and advisor guidance as the project evolves."
    )

    # Sustainability & LCA
    doc.paragraphs[66].text = (
        "Evaluate environmental impact using a simplified LCA mindset. "
        + src.paragraphs[68].text.replace("Sustainability: ", "").strip()
    )
    doc.paragraphs[67].text = src.paragraphs[69].text
    doc.paragraphs[68].text = src.paragraphs[70].text
    doc.paragraphs[69].text = (
        src.paragraphs[71].text + " " + src.paragraphs[72].text + " " + src.paragraphs[73].text
    ).strip()

    # Economic feasibility
    doc.paragraphs[71].text = fix_tex_escapes(src.paragraphs[75].text)
    # 72: şablondaki “The analysis should consider, where applicable:” cümlesi aynı kalır
    doc.paragraphs[73].text = fix_tex_escapes(src.paragraphs[76].text)
    doc.paragraphs[74].text = fix_tex_escapes(src.paragraphs[77].text)
    doc.paragraphs[75].text = fix_tex_escapes(src.paragraphs[78].text).replace("\\0", "0")
    doc.paragraphs[76].text = fix_tex_escapes(src.paragraphs[79].text)
    doc.paragraphs[77].text = fix_tex_escapes(src.paragraphs[80].text)
    doc.paragraphs[78].text = fix_tex_escapes(src.paragraphs[81].text)
    doc.paragraphs[79].text = (
        "A brief discussion on feasibility: the marginal cash cost per analysis session is effectively zero for a static web "
        "prototype, while the benefit is improved interpretability of location uncertainty relative to point-only displays."
    )

    # Risk (şablonda 81 başlık; 82–86 liste)
    for i, j in zip(range(82, 87), range(83, 88)):
        doc.paragraphs[i].text = src.paragraphs[j].text

    # Schedule
    doc.paragraphs[89].text = (
        "Provide a Gantt chart for the project. Below is an indicative weekly plan and task leadership for the ENS 491 track."
    )
    insert_gantt(doc.paragraphs[93], GANTT, "Figure: Summary Gantt chart for ENS 491 (task leaders and weeks).")

    # Ethical
    doc.paragraphs[95].text = (
        "In this section we identify ethical aspects of our project/solution: location traces are sensitive; uncertainty UIs "
        "may be misused to imply rigor; exploratory design may use AI-assisted ideas; movement visualization has dual-use "
        "risks; broader user studies will follow university ethics procedures."
    )
    doc.paragraphs[96].text = src.paragraphs[93].text
    doc.paragraphs[97].text = src.paragraphs[94].text
    doc.paragraphs[98].text = src.paragraphs[95].text
    doc.paragraphs[99].text = src.paragraphs[96].text
    doc.paragraphs[100].text = src.paragraphs[97].text
    doc.paragraphs[102].text = (
        "If there aren’t any, you need to still keep this section and state that there are no ethical issues. "
        "We retain this section because location and uncertainty communication always carry residual ethics considerations."
    )

    # References
    refs = [src.paragraphs[i].text for i in range(99, 109)]
    doc.paragraphs[105].text = "\n".join(refs)

    calibri12_doc(doc)
    doc.save(str(OUT))


if __name__ == "__main__":
    main()
